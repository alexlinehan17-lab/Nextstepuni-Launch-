"""
Generate NextStepUni email draft as a Word document.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING

doc = Document()

# --- Margins: 2.5 cm on all sides ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Helper: apply Calibri 11pt to a run
def style_run(run, bold=False):
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold

# Helper: set single spacing, no extra space before/after
def style_paragraph(paragraph):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

# --- Subject line (bold) ---
p = doc.add_paragraph()
style_paragraph(p)
run_label = p.add_run("Subject: ")
style_run(run_label, bold=True)
run_subject = p.add_run("NextStepUni \u2014 Year 1 Progress & Year 2 Evolution")
style_run(run_subject, bold=True)

# Blank line after subject
blank = doc.add_paragraph()
style_paragraph(blank)

# --- Greeting ---
p = doc.add_paragraph()
style_paragraph(p)
run = p.add_run("Hi Doone, Clodagh,")
style_run(run)

# Blank line
blank = doc.add_paragraph()
style_paragraph(blank)

# --- Body paragraph 1 ---
p = doc.add_paragraph()
style_paragraph(p)
run = p.add_run(
    "Hope you\u2019re both well."
)
style_run(run)

# Blank line
blank = doc.add_paragraph()
style_paragraph(blank)

# --- Body paragraph 2 ---
p = doc.add_paragraph()
style_paragraph(p)
run = p.add_run(
    "We\u2019re about halfway through Year 1 of the NextStepUni programme with the "
    "NEIC schools and things are going well \u2014 all six schools are actively engaged "
    "and the feedback from guidance counsellors has been really valuable."
)
style_run(run)

# Blank line
blank = doc.add_paragraph()
style_paragraph(blank)

# --- Body paragraph 3 ---
p = doc.add_paragraph()
style_paragraph(p)
run = p.add_run(
    "Off the back of what we\u2019ve learned this year and direct input from the schools, "
    "I\u2019ve been developing the next evolution of the programme for Year 2. It builds on "
    "everything that\u2019s working and adds a layer of interactivity and personalisation "
    "that the guidance counsellors have specifically flagged as being impactful for "
    "these students."
)
style_run(run)

# Blank line
blank = doc.add_paragraph()
style_paragraph(blank)

# --- Body paragraph 4 ---
p = doc.add_paragraph()
style_paragraph(p)
run = p.add_run(
    "I\u2019d love to walk you both through where we are and what Year 2 could look like. "
    "I\u2019ve put together a short deck and a supporting document \u2014 should take about "
    "30 minutes."
)
style_run(run)

# Blank line
blank = doc.add_paragraph()
style_paragraph(blank)

# --- Body paragraph 5 ---
p = doc.add_paragraph()
style_paragraph(p)
run = p.add_run(
    "Would either of you have time in the next couple of weeks? Happy to work "
    "around your schedules."
)
style_run(run)

# Blank line
blank = doc.add_paragraph()
style_paragraph(blank)

# --- Sign-off ---
p = doc.add_paragraph()
style_paragraph(p)
run = p.add_run("Thanks,")
style_run(run)

# "Alex" on the next line (no blank line between Thanks and Alex)
p = doc.add_paragraph()
style_paragraph(p)
run = p.add_run("Alex")
style_run(run)

# --- Save ---
output_path = "/Users/alexlinehan/Downloads/NextStepUni_Email_Draft.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
