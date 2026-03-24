"""
Generate NextStepUni Year 2 One-Page Executive Summary
for PwC stakeholders (Doone O'Doherty, Clodagh Dunleavy).
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour constants ──
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x1A)
BODY_GREY = RGBColor(0x33, 0x33, 0x33)
SUBTITLE_GREY = RGBColor(0x66, 0x66, 0x66)
TERRACOTTA = RGBColor(0xCC, 0x78, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF9, 0xF5, 0xF3)
LINE_GREY = RGBColor(0xCC, 0xCC, 0xCC)

FONT_NAME = "Calibri"
OUTPUT_PATH = "/Users/alexlinehan/Downloads/NextStepUni_Year2_OnePager.docx"


def set_font(run, size_pt, color, bold=False, font_name=FONT_NAME):
    """Apply font formatting to a run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.bold = bold
    # Ensure Calibri is used for East Asian text too
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rPr.insert(0, rFonts)


def set_paragraph_spacing(paragraph, before_pt=0, after_pt=0, line_spacing=1.0):
    """Set paragraph spacing tightly."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = line_spacing


def add_section_header(doc, text, before_pt=6, after_pt=2):
    """Add a terracotta uppercase section header."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=before_pt, after_pt=after_pt)
    run = p.add_run(text.upper())
    set_font(run, 10, TERRACOTTA, bold=True)
    return p


def add_body_paragraph(doc, text, before_pt=0, after_pt=2, line_spacing=1.0):
    """Add a body text paragraph."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=before_pt, after_pt=after_pt, line_spacing=line_spacing)
    run = p.add_run(text)
    set_font(run, 9.5, BODY_GREY)
    return p


def add_bullet(doc, text, before_pt=0, after_pt=1):
    """Add a compact bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    # Clear default text and add our own run
    p.clear()
    run = p.add_run(text)
    set_font(run, 9.5, BODY_GREY)
    set_paragraph_spacing(p, before_pt=before_pt, after_pt=after_pt, line_spacing=1.0)
    # Tighten bullet indent
    pf = p.paragraph_format
    pf.left_indent = Cm(0.6)
    pf.first_line_indent = Cm(-0.3)
    return p


def add_thin_line(doc):
    """Add a thin horizontal line via a bottom-bordered paragraph."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=2, after_pt=2)
    # Add bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def shade_cell(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, size_pt, color, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    set_paragraph_spacing(p, before_pt=1, after_pt=1, line_spacing=1.0)
    run = p.add_run(text)
    set_font(run, size_pt, color, bold=bold)


def remove_cell_borders(cell, sides=None):
    """Remove specific borders from a cell or set them to a subtle style."""
    pass  # We'll handle borders at the table level


def main():
    doc = Document()

    # ── Page setup: 2cm margins ──
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Remove headers/footers page numbers
    section.different_first_page_header_footer = False
    header = section.header
    header.is_linked_to_previous = True

    # ── Set default document font ──
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(9.5)
    style.font.color.rgb = BODY_GREY
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.0

    # ── HEADER AREA ──
    # Title
    title_p = doc.add_paragraph()
    set_paragraph_spacing(title_p, before_pt=0, after_pt=0)
    title_run = title_p.add_run("NextStepUni")
    set_font(title_run, 18, DARK_CHARCOAL, bold=True)
    title_run2 = title_p.add_run(" \u2014 Year 2 Proposal")
    set_font(title_run2, 18, DARK_CHARCOAL, bold=True)

    # Subtitle
    sub_p = doc.add_paragraph()
    set_paragraph_spacing(sub_p, before_pt=0, after_pt=1)
    sub_run = sub_p.add_run("PwC Empowering Futures Programme")
    set_font(sub_run, 12, SUBTITLE_GREY)

    # Thin line
    add_thin_line(doc)

    # ── SECTION 1: PROGRAMME SUMMARY ──
    add_section_header(doc, "Programme Summary", before_pt=4, after_pt=1)
    add_body_paragraph(
        doc,
        "NextStepUni is a study skills and exam preparation programme for Leaving Cert students "
        "in DEIS schools. Funded by PwC under the Empowering Futures initiative, it currently "
        "operates across 6 NEIC schools reaching ~500 students through in-person workshops and "
        "an online learning platform.",
        after_pt=2,
    )

    # ── SECTION 2: YEAR 1 PROGRESS ──
    add_section_header(doc, "Year 1 Progress", before_pt=4, after_pt=1)
    bullets_y1 = [
        "Programme live across all 6 NEIC schools since October 2025",
        "~500 students onboarded (5th and 6th years)",
        "In-person workshops delivered in all schools",
        "Guidance counsellor feedback collected and incorporated into Year 2 planning",
    ]
    for b in bullets_y1:
        add_bullet(doc, b)

    # ── SECTION 3: YEAR 2 EVOLUTION ──
    add_section_header(doc, "Year 2 Evolution", before_pt=4, after_pt=1)
    add_body_paragraph(
        doc,
        "Based on Year 1 learnings and direct feedback from guidance counsellors, the online "
        "component is evolving from a passive video course into a fully interactive learning "
        "platform. The platform includes 52 research-backed interactive modules, personalised "
        "student onboarding, built-in study tools (flashcards, study planner, exam simulators), "
        "AI-powered study support, and gamification features that drive consistent engagement. "
        "In-person workshops continue unchanged.",
        after_pt=2,
    )

    # ── SECTION 4: KEY DIFFERENTIATORS ──
    add_section_header(doc, "Key Differentiators", before_pt=4, after_pt=1)
    bullets_diff = [
        "Built from guidance counsellor feedback \u2014 not generic, built for these students",
        "Research-backed \u2014 every module grounded in educational psychology",
        "Measurable impact \u2014 real-time tracking of student engagement, completion, and progress",
        "Socioeconomically conscious \u2014 no assumptions about resources; works on any device",
    ]
    for b in bullets_diff:
        add_bullet(doc, b)

    # ── SECTION 5: INVESTMENT ──
    add_section_header(doc, "Investment", before_pt=4, after_pt=2)

    # Create a clean 3-column, 4-row table
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    # Set preferred table width to about 60% of page width for a compact look
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="4200" w:type="pct"/>')
    # Remove existing tblW if any
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblPr.append(tblW)

    # Remove default table borders, then add subtle ones
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="2" w:space="0" w:color="DDDDDD"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="2" w:space="0" w:color="DDDDDD"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="EEEEEE"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    # Cell margins for compactness
    tblCellMar = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="20" w:type="dxa"/>'
        f'  <w:bottom w:w="20" w:type="dxa"/>'
        f'  <w:left w:w="80" w:type="dxa"/>'
        f'  <w:right w:w="80" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    for existing in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(existing)
    tblPr.append(tblCellMar)

    # Header row
    header_data = ["", "Year 1", "Year 2"]
    for i, txt in enumerate(header_data):
        cell = table.cell(0, i)
        shade_cell(cell, "CC785C")
        align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        set_cell_text(cell, txt, 9, WHITE, bold=True, alignment=align)

    # Data rows
    rows_data = [
        ("Schools", "6 NEIC", "6 NEIC"),
        ("Students", "~500", "~500\u2013600"),
        ("Annual investment", "\u20ac18,000", "\u20ac30,000"),
    ]
    for row_idx, (label, y1, y2) in enumerate(rows_data, start=1):
        set_cell_text(table.cell(row_idx, 0), label, 9, BODY_GREY, bold=True)
        set_cell_text(table.cell(row_idx, 1), y1, 9, BODY_GREY, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(row_idx, 2), y2, 9, BODY_GREY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        # Alternate row shading
        if row_idx % 2 == 1:
            for c in range(3):
                shade_cell(table.cell(row_idx, c), "F9F5F3")

    # Investment note
    note_p = doc.add_paragraph()
    set_paragraph_spacing(note_p, before_pt=2, after_pt=1, line_spacing=1.0)
    note_run = note_p.add_run(
        "The \u20ac12,000 increase covers platform infrastructure, AI tools, and ongoing maintenance. "
        "Platform development was completed at no cost to PwC (equivalent external development: "
        "\u20ac40,000\u201360,000)."
    )
    set_font(note_run, 8.5, SUBTITLE_GREY, bold=False)

    # ── SECTION 6: NEXT STEPS ──
    add_section_header(doc, "Next Steps", before_pt=4, after_pt=1)
    bullets_next = [
        "Agree Year 2 delivery model and investment",
        "Year 1 review \u2014 June 2026",
        "Year 2 launch \u2014 October 2026",
    ]
    for b in bullets_next:
        add_bullet(doc, b)

    # ── FOOTER ──
    add_thin_line(doc)
    footer_p = doc.add_paragraph()
    set_paragraph_spacing(footer_p, before_pt=2, after_pt=0)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("Alex Linehan  |  NextStepUni  |  PwC Empowering Futures")
    set_font(footer_run, 8.5, SUBTITLE_GREY)

    # ── Save ──
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    main()
