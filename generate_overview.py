#!/usr/bin/env python3
"""
Generate the NextStepUni Programme Overview & Evidence Base document
for the PwC Empowering Futures meeting.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ──────────────────────────────────────────────────────────────
# Constants
# ──────────────────────────────────────────────────────────────

FONT_NAME = "Calibri"
ACCENT_COLOR = RGBColor(0x2D, 0x37, 0x48)  # Dark navy/charcoal
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
BODY_GRAY = RGBColor(0x2A, 0x2A, 0x2A)
LIGHT_LINE = RGBColor(0xBB, 0xBB, 0xBB)

OUTPUT_PATH = "/Users/alexlinehan/Downloads/NextStepUni_Programme_Overview.docx"


def set_font(run, size=11, bold=False, color=None, italic=False, name=FONT_NAME):
    """Helper to set font properties on a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Force Calibri for East Asian text too
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{name}" w:cs="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:eastAsia'), name)
        rFonts.set(qn('w:cs'), name)


def add_paragraph(doc, text, size=11, bold=False, color=None, alignment=None,
                  space_before=0, space_after=6, line_spacing=1.15, italic=False):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color or BODY_GRAY, italic=italic)
    return p


def add_body_text(doc, text, space_before=0, space_after=8):
    """Add standard body text."""
    return add_paragraph(doc, text, size=11, color=BODY_GRAY,
                         space_before=space_before, space_after=space_after,
                         line_spacing=1.15)


def add_section_header(doc, text, space_before=18):
    """Add a section header (16pt bold, accent colour)."""
    return add_paragraph(doc, text, size=16, bold=True, color=ACCENT_COLOR,
                         space_before=space_before, space_after=8,
                         line_spacing=1.0)


def add_subsection_header(doc, text, space_before=14):
    """Add a subsection header (13pt bold, accent colour)."""
    return add_paragraph(doc, text, size=13, bold=True, color=ACCENT_COLOR,
                         space_before=space_before, space_after=6,
                         line_spacing=1.0)


def add_bullet(doc, text, size=11, bold_prefix=None, indent_level=0, space_after=4):
    """Add a bullet-point paragraph."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15

    # Indent
    left_indent = Cm(1.2 + (indent_level * 0.8))
    hanging = Cm(0.5)
    pf.left_indent = left_indent
    pf.first_line_indent = -hanging

    if bold_prefix:
        bullet_run = p.add_run("\u2022  ")
        set_font(bullet_run, size=size, color=BODY_GRAY)
        bold_run = p.add_run(bold_prefix)
        set_font(bold_run, size=size, bold=True, color=BODY_GRAY)
        rest_run = p.add_run(text)
        set_font(rest_run, size=size, color=BODY_GRAY)
    else:
        run = p.add_run(f"\u2022  {text}")
        set_font(run, size=size, color=BODY_GRAY)
    return p


def add_horizontal_line(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # Use bottom border on the paragraph
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="AAAAAA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_page_break(doc):
    """Insert a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)


def add_page_numbers(doc):
    """Add page numbers to the footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        # Add page number field
        run = p.add_run()
        set_font(run, size=9, color=RGBColor(0x99, 0x99, 0x99))
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)

        run2 = p.add_run()
        set_font(run2, size=9, color=RGBColor(0x99, 0x99, 0x99))
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)

        run3 = p.add_run()
        set_font(run3, size=9, color=RGBColor(0x99, 0x99, 0x99))
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)


# ──────────────────────────────────────────────────────────────
# Build the document
# ──────────────────────────────────────────────────────────────

import docx.enum.text

doc = Document()

# -- Page setup: 2.5 cm margins --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Set default font --
style = doc.styles['Normal']
font = style.font
font.name = FONT_NAME
font.size = Pt(11)
font.color.rgb = BODY_GRAY
rFonts = style.element.rPr.find(qn('w:rFonts')) if style.element.rPr is not None else None
if style.element.rPr is None:
    rPr = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/></w:rPr>')
    style.element.append(rPr)

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════

# Spacer
add_paragraph(doc, "", size=11, space_before=0, space_after=0)
add_paragraph(doc, "", size=11, space_before=0, space_after=0)
add_paragraph(doc, "", size=11, space_before=0, space_after=0)

# Title
add_paragraph(doc, "NextStepUni", size=28, bold=True, color=ACCENT_COLOR,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=60, space_after=4,
              line_spacing=1.0)

add_paragraph(doc, "Programme Overview & Evidence Base", size=18, bold=False,
              color=ACCENT_COLOR, alignment=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=0, space_after=4, line_spacing=1.0)

add_paragraph(doc, "Year 2 Evolution: Interactive Learning Platform", size=14,
              bold=False, color=RGBColor(0x55, 0x55, 0x55),
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=20,
              line_spacing=1.0)

# Horizontal line
add_horizontal_line(doc)

add_paragraph(doc, "Prepared for PwC Empowering Futures Programme", size=12,
              bold=False, color=RGBColor(0x66, 0x66, 0x66),
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=4,
              line_spacing=1.0)

add_paragraph(doc, "February 2026", size=12, bold=False,
              color=RGBColor(0x66, 0x66, 0x66),
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0,
              line_spacing=1.0)

# Page break
add_page_break(doc)


# ══════════════════════════════════════════════════════════════
# SECTION 1: PROGRAMME OVERVIEW
# ══════════════════════════════════════════════════════════════

add_section_header(doc, "1.  Programme Overview", space_before=0)

add_body_text(doc, (
    "NextStepUni is a study skills and exam preparation programme designed specifically for "
    "Leaving Certificate students in DEIS schools. The programme equips 16\u201318 year olds with "
    "the evidence-based study strategies, mindset tools, and exam techniques they need to "
    "achieve their academic potential and access higher education. At its core, NextStepUni "
    "exists to close the preparation gap \u2014 ensuring that students from disadvantaged "
    "backgrounds have the same quality of study support that is routinely available to their "
    "peers in better-resourced schools."
))

add_body_text(doc, (
    "In Year 1, the programme was delivered across six DEIS schools in Dublin\u2019s North East "
    "Inner City, reaching approximately 500 students through a combination of in-person "
    "workshops and an online video course. The workshops provided direct engagement with "
    "students and teachers, while the online course offered a library of study skills content "
    "that students could access in their own time. This first year established the programme\u2019s "
    "presence in schools, built strong relationships with guidance counsellors, and generated "
    "valuable feedback on how the online component could be improved."
))

add_body_text(doc, (
    "Year 2 represents the natural evolution of that foundation. Based on detailed feedback "
    "from guidance counsellors and grounded in educational psychology research, the online "
    "component is being rebuilt as a fully interactive learning platform. The passive video "
    "course is replaced by 52 structured, interactive modules that ask students to engage, "
    "reflect, plan, and apply \u2014 rather than simply watch and absorb. The result is a "
    "programme that is more effective, more measurable, and better suited to how students "
    "actually learn."
))


# ══════════════════════════════════════════════════════════════
# SECTION 2: WHAT THE PLATFORM INCLUDES
# ══════════════════════════════════════════════════════════════

add_section_header(doc, "2.  What the Platform Includes")

add_body_text(doc, (
    "The Year 2 platform is a purpose-built React web application, designed for mobile-first "
    "access on any device with an internet connection. Every element of the platform has been "
    "designed with the target cohort in mind: clear language, no assumptions about prior "
    "knowledge, and a structured pathway that builds understanding step by step."
))

# --- 2.1 ---
add_subsection_header(doc, "2.1  52 Interactive Learning Modules")

add_body_text(doc, (
    "The platform delivers 52 learning modules organised across five categories, each "
    "addressing a different dimension of academic success."
))

# Category bullets with bold prefix
add_bullet(doc, bold_prefix="Architecture of Your Mind (13 modules) \u2014 ",
           text="Mindset, motivation, goal-setting, and resilience. Topics include growth mindset, "
                "the science of grit, self-efficacy, emotional intelligence, and future planning.")

add_bullet(doc, bold_prefix="The Science of Growth (10 modules) \u2014 ",
           text="The neuroscience of learning, neuroplasticity, the power of mistakes, effective "
                "struggle, and how praise and feedback shape performance.")

add_bullet(doc, bold_prefix="Learning Cheat Codes (15 modules) \u2014 ",
           text="Evidence-based study techniques including active recall, spaced repetition, "
                "interleaving, elaborative interrogation, the note-taking paradox, cognitive load "
                "management, and more.")

add_bullet(doc, bold_prefix="Subject-Specific Strategies (8 modules) \u2014 ",
           text="Tailored approaches for Maths, English, Sciences, Languages, Business, Humanities, "
                "Creative subjects, and Applied/Technical subjects. Each module breaks down exactly "
                "how that subject\u2019s exam works and how to maximise marks.")

add_bullet(doc, bold_prefix="The Exam Zone (6 modules) \u2014 ",
           text="Exam-specific preparation covering Leaving Cert strategy, reverse-engineering study "
                "plans, exam hall techniques, crisis management, game day routines, and points "
                "optimisation.")

add_body_text(doc, (
    "Each module is built around a set of interactive features designed to move students from "
    "passive consumption to active engagement:"
), space_before=8)

add_bullet(doc, bold_prefix="Interactive Highlight Tooltips \u2014 ",
           text="Key concepts are highlighted throughout each module. When tapped, a tooltip "
                "explains the concept in plain, accessible language, reinforcing understanding "
                "without breaking the reading flow.")

add_bullet(doc, bold_prefix="Micro-Commitments \u2014 ",
           text="At the end of each section, students are prompted to make a small, specific, "
                "actionable commitment. This transforms learning from abstract to personal.")

add_bullet(doc, bold_prefix="Personal Stories \u2014 ",
           text="First-person accounts from relatable students are woven into modules to build a "
                "sense of belonging and normalise the challenges of studying.")

add_bullet(doc, bold_prefix="Progressive Unlock \u2014 ",
           text="Students work through sections sequentially, building understanding step by step. "
                "This prevents skipping ahead and ensures foundational concepts are absorbed before "
                "more advanced material.")

add_bullet(doc, bold_prefix="Embedded Interactive Tools \u2014 ",
           text="Modules include built-in tools such as WOOP goal planners, exam triage simulators, "
                "grade calculators, breathing exercises, and study environment planners \u2014 turning "
                "theory into immediate practice.")

# --- 2.2 ---
add_subsection_header(doc, "2.2  Personalised Onboarding")

add_body_text(doc, (
    "When students first log in, they complete a two-minute onboarding flow that captures their "
    "Leaving Certificate subjects, current and target grades for each subject, Higher or Ordinary "
    "level selections, exam date, and preferred study days. This information creates a personalised "
    "study profile for each student."
))

add_body_text(doc, (
    "The platform uses this data to calculate each student\u2019s current projected CAO points and "
    "their target points, using the standard best-six CAO scoring formula (capped at 625 points). "
    "This gives every student a clear, personalised picture of where they currently stand and "
    "where they want to get to \u2014 a powerful motivational anchor that frames their entire "
    "experience on the platform."
))

# --- 2.3 ---
add_subsection_header(doc, "2.3  Innovation Zone (AI-Powered Study Tools)")

add_body_text(doc, (
    "The Innovation Zone is a suite of AI-powered study tools, built on Google Gemini, that give "
    "students access to on-demand academic support. The tools include:"
))

add_bullet(doc, bold_prefix="AI Study Assistant \u2014 ",
           text="Students can ask questions about any topic and receive clear, Leaving Cert-relevant "
                "explanations tailored to the Irish curriculum.")

add_bullet(doc, bold_prefix="Essay Feedback Tool \u2014 ",
           text="Students paste in draft essays and receive structured, constructive feedback on "
                "content, structure, and technique.")

add_bullet(doc, bold_prefix="Subject-Specific Help \u2014 ",
           text="Tailored support across all Leaving Cert subjects, aligned with the specific "
                "requirements of the Irish syllabus and exam format.")

add_body_text(doc, (
    "These tools are designed to supplement, not replace, teacher support. For many students in "
    "disadvantaged areas, private tutoring is not an option. The Innovation Zone provides a "
    "level of on-demand, personalised academic help that was previously only available to "
    "students whose families could afford it."
), space_before=4)

# --- 2.4 ---
add_subsection_header(doc, "2.4  Progress Tracking & Impact Data")

add_body_text(doc, (
    "The platform tracks detailed engagement data for every student, including which modules "
    "have been completed, how far they have progressed through each module, their interactive "
    "responses (goal plans, self-assessments, study schedules), and overall engagement patterns "
    "over time."
))

add_body_text(doc, (
    "This provides both NextStepUni and PwC with granular, real-time data on programme impact "
    "\u2014 the first time this level of reporting has been available for the programme. Year 1 "
    "could measure whether a student watched a video; Year 2 can measure whether they engaged "
    "with the material, set goals, made commitments, and returned over time."
))


# ══════════════════════════════════════════════════════════════
# SECTION 3: THE EVIDENCE BASE
# ══════════════════════════════════════════════════════════════

add_section_header(doc, "3.  The Evidence Base")

add_body_text(doc, (
    "The shift from a passive video course to an interactive learning platform is not a stylistic "
    "preference \u2014 it is grounded in decades of research in educational psychology and the "
    "learning sciences. The following sections summarise the key evidence behind each element "
    "of the platform\u2019s design."
))

# --- 3.1 ---
add_subsection_header(doc, "3.1  Active Learning vs. Passive Delivery")

add_body_text(doc, (
    "The Year 1 video course follows a passive delivery model: students watch content, then "
    "complete a short quiz. While this approach has value as an introductory format, research "
    "consistently shows it is one of the least effective ways to produce lasting learning."
))

add_body_text(doc, (
    "Freeman et al. (2014) conducted a landmark meta-analysis of 225 studies across STEM "
    "disciplines and found that active learning increases exam performance by half a grade on "
    "average. Students in traditional passive lecture courses were 1.5 times more likely to fail "
    "than those in active learning environments. This finding has been described as one of the "
    "largest and most robust effects in educational research."
))

add_body_text(doc, (
    "Chi and Wylie\u2019s (2014) ICAP framework provides a useful lens for understanding this. "
    "Learning outcomes improve progressively as students move from Passive engagement (watching "
    "or listening) to Active engagement (doing something with the material) to Constructive "
    "engagement (creating new understanding) to Interactive engagement (discussing and "
    "collaborating). The Year 1 course sits primarily in the Passive category. The Year 2 "
    "platform moves students firmly into the Constructive and Interactive categories through "
    "its tooltips, micro-commitments, embedded tools, and AI-powered feedback."
))

add_body_text(doc, (
    "Prince (2004) reinforces this in a comprehensive review of the active learning literature, "
    "confirming that active learning methods produce superior outcomes across disciplines, "
    "student populations, and assessment types."
))

# --- 3.2 ---
add_subsection_header(doc, "3.2  Spaced Repetition & Active Recall")

add_body_text(doc, (
    "Two of the most robustly supported strategies in cognitive psychology are active recall "
    "(retrieving information from memory rather than re-reading it) and spaced repetition "
    "(distributing study over time rather than cramming)."
))

add_body_text(doc, (
    "Roediger and Butler (2011) demonstrated that testing yourself on material is significantly "
    "more effective than re-reading or re-watching \u2014 even when the total study time is the same. "
    "This \u201ctesting effect\u201d is one of the most replicated findings in the learning sciences. The "
    "platform\u2019s interactive elements \u2014 micro-commitments, self-assessments, and embedded "
    "tools \u2014 function as retrieval practice, prompting students to actively recall and apply "
    "what they have learned."
))

add_body_text(doc, (
    "Cepeda et al. (2006) showed that spacing study sessions over time produces 10\u201330% better "
    "long-term retention compared to massed practice (cramming). The platform\u2019s modular "
    "structure and progressive unlock system naturally encourage spaced engagement, as students "
    "return to the platform over days and weeks rather than consuming everything in a single "
    "sitting."
))

add_body_text(doc, (
    "Dunlosky et al. (2013) reviewed ten common learning strategies and rated practice testing "
    "and distributed practice as the two highest-utility techniques \u2014 effective across a wide "
    "range of learners, materials, and contexts. Both strategies are embedded directly into the "
    "platform\u2019s design."
))

# --- 3.3 ---
add_subsection_header(doc, "3.3  Interleaving & Elaborative Interrogation")

add_body_text(doc, (
    "Rohrer and Taylor (2007) found that interleaving \u2014 mixing different topics or problem "
    "types during study \u2014 produces significantly better long-term retention than blocking "
    "(studying one topic exhaustively before moving to the next). The platform\u2019s organisation "
    "across five distinct module categories naturally encourages students to interleave their "
    "learning, moving between mindset, technique, and subject-specific content."
))

add_body_text(doc, (
    "Elaborative interrogation \u2014 the practice of asking \u201cwhy\u201d and \u201chow\u201d questions about "
    "material being studied \u2014 was rated as a moderate-to-high utility strategy by Dunlosky "
    "et al. (2013). The platform\u2019s Interactive Highlight Tooltips prompt exactly this kind "
    "of deeper processing: when students encounter a highlighted concept and tap to learn more, "
    "they are engaging in a form of self-explanation that strengthens understanding and "
    "retention."
))

# --- 3.4 ---
add_subsection_header(doc, "3.4  Interactive & Personalised Learning")

add_body_text(doc, (
    "Pashler et al. (2007) found that personalising learning pathways \u2014 adapting content "
    "and pacing to individual learner profiles \u2014 improves both engagement and outcomes. This "
    "effect is particularly pronounced for students who are at risk of disengaging from "
    "education, a group that includes many of the students NextStepUni serves."
))

add_body_text(doc, (
    "A US Department of Education meta-analysis by Means et al. (2013) found that online "
    "learning environments incorporating interactive elements outperform both pure face-to-face "
    "instruction and purely passive online delivery. The combination of interactivity with the "
    "flexibility of online access produces the strongest outcomes."
))

add_body_text(doc, (
    "VanLehn (2011) demonstrated that interactive problem-solving produces learning gains "
    "comparable to one-on-one human tutoring, and significantly better than passive instruction. "
    "The platform\u2019s embedded tools, AI-powered feedback, and structured reflection exercises "
    "are designed to approximate this kind of interactive, responsive learning experience."
))

# --- 3.5 ---
add_subsection_header(doc, "3.5  The Case for DEIS Students Specifically")

add_body_text(doc, (
    "The evidence above applies broadly, but it carries particular weight for students from "
    "disadvantaged backgrounds. These students are disproportionately affected by passive "
    "delivery models because they often lack the supplementary support that compensates for "
    "passive instruction elsewhere \u2014 private tutors, structured study groups, well-resourced "
    "home environments, and parents who can provide academic guidance."
))

add_body_text(doc, (
    "Interactive platforms provide a levelling effect. By embedding tools, feedback, and "
    "structured support directly into the learning experience, the platform gives every student "
    "access to high-quality study support regardless of their home situation. A student "
    "completing a WOOP goal-planning exercise on the platform at 10pm on a Tuesday night is "
    "receiving the same quality of structured guidance that a student with a private tutor "
    "might get in a paid session."
))

add_body_text(doc, (
    "The platform was built with this socioeconomic awareness at every level. It makes no "
    "assumptions about access to private tutors, expensive equipment, or quiet study spaces. "
    "Every tool and every module is fully accessible on any device with an internet connection, "
    "including a smartphone. This is not an afterthought \u2014 it is a core design principle."
))


# ══════════════════════════════════════════════════════════════
# SECTION 4: YEAR 2 DELIVERY MODEL
# ══════════════════════════════════════════════════════════════

add_section_header(doc, "4.  Year 2 Delivery Model")

add_body_text(doc, (
    "Year 2 of the programme runs from October 2026 to June 2027, covering the full academic "
    "year for both 5th Year and 6th Year students."
))

add_bullet(doc, text="Same six DEIS schools in Dublin\u2019s North East Inner City, reaching "
                     "approximately 500\u2013600 students across 5th and 6th Year.")

add_bullet(doc, text="In-person workshops continue as in Year 1, providing direct engagement "
                     "with students and teachers in each school.")

add_bullet(doc, text="The online component transitions from the passive video course to the "
                     "fully interactive learning platform described in this document.")

add_bullet(doc, text="Students are onboarded via the personalised setup flow, establishing "
                     "their subject profiles, grade targets, and projected CAO points from "
                     "day one.")

add_bullet(doc, text="Ongoing support is available through the Innovation Zone AI tools, giving "
                     "students on-demand academic help outside school hours.")

add_bullet(doc, text="A continuous feedback loop with guidance counsellors ensures the programme "
                     "remains responsive to student needs throughout the year.")


# ══════════════════════════════════════════════════════════════
# SECTION 5: INVESTMENT SUMMARY
# ══════════════════════════════════════════════════════════════

add_section_header(doc, "5.  Investment Summary")

add_body_text(doc, "The total investment for Year 2 of the programme is \u20ac30,000, broken down as follows:")

# Investment table
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'

# Set table width
for cell in table.columns[0].cells:
    cell.width = Cm(12)
for cell in table.columns[1].cells:
    cell.width = Cm(4)

data = [
    ("Programme delivery (workshops + platform access)", "\u20ac18,000"),
    ("Platform infrastructure (hosting, database, authentication)", "\u20ac3,000"),
    ("AI-powered study tools (Google Gemini API)", "\u20ac3,500"),
    ("Platform maintenance, support & content updates", "\u20ac5,500"),
]

# Header row
header_cells = table.rows[0].cells
for i, text in enumerate(["Item", "Cost"]):
    header_cells[i].text = ""
    p = header_cells[i].paragraphs[0]
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
    # Dark background
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2D3748" w:val="clear"/>')
    header_cells[i]._element.get_or_add_tcPr().append(shading)

# Data rows
for row_idx, (item, cost) in enumerate(data):
    cells = table.rows[row_idx + 1].cells
    cells[0].text = ""
    p0 = cells[0].paragraphs[0]
    r0 = p0.add_run(item)
    set_font(r0, size=11, color=BODY_GRAY)

    cells[1].text = ""
    p1 = cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run(cost)
    set_font(r1, size=11, color=BODY_GRAY)

    # Alternate row shading
    if row_idx % 2 == 0:
        for cell in cells:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F7F7F7" w:val="clear"/>')
            cell._element.get_or_add_tcPr().append(shading)

# Total row
total_cells = table.rows[5].cells
total_cells[0].text = ""
p_total = total_cells[0].paragraphs[0]
r_total = p_total.add_run("Total")
set_font(r_total, size=11, bold=True, color=ACCENT_COLOR)

total_cells[1].text = ""
p_total_val = total_cells[1].paragraphs[0]
p_total_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_total_val = p_total_val.add_run("\u20ac30,000")
set_font(r_total_val, size=11, bold=True, color=ACCENT_COLOR)

# Bold border on total row
for cell in total_cells:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EDF2F7" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)

# Set cell padding for all cells
for row in table.rows:
    for cell in row.cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="60" w:type="dxa"/>'
            f'  <w:left w:w="100" w:type="dxa"/>'
            f'  <w:bottom w:w="60" w:type="dxa"/>'
            f'  <w:right w:w="100" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

add_body_text(doc, "", space_before=8, space_after=4)

add_body_text(doc, (
    "It is worth noting that the interactive learning platform itself was designed and developed "
    "by Alex Linehan at no cost to PwC. An equivalent platform, if commissioned from an external "
    "development agency, would typically cost in the range of \u20ac40,000\u2013\u20ac60,000. The \u20ac12,000 "
    "increase from Year 1 to Year 2 covers the ongoing running costs of the platform only: "
    "hosting, database infrastructure, AI API usage, and maintenance."
))


# ──────────────────────────────────────────────────────────────
# Page numbers
# ──────────────────────────────────────────────────────────────

add_page_numbers(doc)

# ──────────────────────────────────────────────────────────────
# Save
# ──────────────────────────────────────────────────────────────

os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes")
